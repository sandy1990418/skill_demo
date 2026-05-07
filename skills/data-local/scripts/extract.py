"""
Extract text and tables from local files (PDF, DOCX, XLSX, CSV, JSON, TXT).

Usage:
    python extract.py --file <path> [--pages 1-5] [--output result.md]

Output: structured markdown to stdout (or --output file).
"""

import argparse
import json
import sys
from pathlib import Path


def extract_pdf(path: Path, pages: str | None) -> str:
    import pdfplumber

    page_range = None
    if pages:
        parts = pages.split("-")
        start = int(parts[0]) - 1
        end = int(parts[1]) if len(parts) > 1 else start + 1
        page_range = range(start, end)

    chunks = []
    with pdfplumber.open(path) as pdf:
        all_pages = [pdf.pages[i] for i in page_range] if page_range else pdf.pages
        for page in all_pages:
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            chunks.append(f"### Page {page.page_number}\n\n{text}")
            for tbl in tables:
                if not tbl:
                    continue
                header = tbl[0]
                rows = tbl[1:]
                md = "| " + " | ".join(str(c or "") for c in header) + " |\n"
                md += "| " + " | ".join("---" for _ in header) + " |\n"
                for row in rows:
                    md += "| " + " | ".join(str(c or "") for c in row) + " |\n"
                chunks.append(md)
    return "\n\n".join(chunks)


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    chunks = []
    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text)
    for i, tbl in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        if not rows:
            continue
        md = f"\n**Table {i+1}**\n\n"
        md += "| " + " | ".join(rows[0]) + " |\n"
        md += "| " + " | ".join("---" for _ in rows[0]) + " |\n"
        for row in rows[1:]:
            md += "| " + " | ".join(row) + " |\n"
        chunks.append(md)
    return "\n\n".join(chunks)


def extract_xlsx(path: Path) -> str:
    import pandas as pd

    sheets = pd.read_excel(path, sheet_name=None)
    chunks = []
    for name, df in sheets.items():
        chunks.append(f"### Sheet: {name} ({df.shape[0]} rows × {df.shape[1]} cols)\n")
        chunks.append(df.to_markdown(index=False))
        chunks.append(f"\n**Stats:**\n{df.describe(include='all').to_markdown()}")
    return "\n\n".join(chunks)


def extract_csv(path: Path) -> str:
    import pandas as pd

    df = pd.read_csv(path)
    return (
        f"### {path.name} ({df.shape[0]} rows × {df.shape[1]} cols)\n\n"
        + df.to_markdown(index=False)
        + f"\n\n**Stats:**\n{df.describe(include='all').to_markdown()}"
    )


def extract_json(path: Path) -> str:
    data = json.loads(path.read_text(encoding="utf-8"))
    return f"```json\n{json.dumps(data, indent=2, ensure_ascii=False)}\n```"


def extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


EXTRACTORS = {
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".xls":  extract_xlsx,
    ".csv":  extract_csv,
    ".json": extract_json,
    ".txt":  extract_txt,
    ".md":   extract_txt,
}


def main():
    parser = argparse.ArgumentParser(description="Extract content from local files")
    parser.add_argument("--file", required=True, help="Path to file (or directory to batch-extract)")
    parser.add_argument("--pages", help="PDF page range, e.g. 1-5")
    parser.add_argument("--output", help="Write result to this file instead of stdout")
    args = parser.parse_args()

    target = Path(args.file)
    results = []

    files = sorted(target.iterdir()) if target.is_dir() else [target]

    for f in files:
        ext = f.suffix.lower()
        if ext not in EXTRACTORS:
            print(f"[skip] {f.name} — unsupported type {ext}", file=sys.stderr)
            continue

        print(f"[extract] {f.name}", file=sys.stderr)
        extractor = EXTRACTORS[ext]

        try:
            if ext == ".pdf":
                content = extractor(f, args.pages)
            else:
                content = extractor(f)
        except Exception as e:
            print(f"[error] {f.name}: {e}", file=sys.stderr)
            continue

        results.append(f"## Source: {f.name}\n\n{content}")

    output = "\n\n---\n\n".join(results)

    if args.output:
        Path(args.output).write_text(output, encoding="utf-8")
        print(f"[done] written to {args.output}", file=sys.stderr)
    else:
        print(output)


if __name__ == "__main__":
    main()
